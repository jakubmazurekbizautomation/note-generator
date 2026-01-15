import streamlit as st
import fitz
from openai import OpenAI
from docx import Document
import io
import json

# Konfiguracja
client = OpenAI(api_key=st.secrets["OPENAI_KEY"])

st.title("📚 Generator Notatek + Quiz")

# Tabs
tab1, tab2 = st.tabs(["📝 Generuj Notatki", "🎯 Quiz Interaktywny"])

# ===== TAB 1: NOTATKI =====
with tab1:
    # Upload
    pdf = st.file_uploader("Wrzuć PDF", type="pdf")

    # Tryb
    tryb = st.radio("Tryb:", ["Przepisz 1:1", "Skrót", "Rozszerzone z wyjaśnieniami"])

    # Custom prompt
    st.markdown("---")
    st.subheader("🎯 Dodatkowe instrukcje (opcjonalnie)")
    custom_prompt = st.text_area(
        "Np: 'Rozwiń temat elektrolitów, to moja słaba strona'",
        placeholder="Wpisz na czym Ci szczególnie zależy...",
        height=100
    )

    # Weryfikacja - tylko dla trybu rozszerzonego
    if tryb == "Rozszerzone z wyjaśnieniami":
        weryfikacja = st.checkbox("✅ Sprawdź i popraw błędy + dodaj dodatkowe wyjaśnienia", value=True)
    else:
        weryfikacja = False

    if pdf and st.button("🚀 Generuj Notatki"):
        with st.spinner("Przetwarzam..."):
            # Wyciągnij tekst
            doc = fitz.open(stream=pdf.read(), filetype="pdf")
            tekst = ""
            for page in doc:
                tekst += page.get_text()
            
            # Zapisz tekst w session_state do quizu
            st.session_state['pdf_tekst'] = tekst
            
            # Prompt bazowy
            prompty = {
                "Przepisz 1:1": """Przepisz dokładnie tekst z tego materiału edukacyjnego do formatu notatek:

✓ Zachowaj CAŁĄ treść bez skracania
✓ Popraw tylko formatowanie (dodaj nagłówki, punktory gdzie pasują)
✓ Popraw ewentualne błędy ortograficzne
✓ NIE zmieniaj treści merytorycznej
✓ NIE dodawaj niczego od siebie

Materiał:""",
                
                "Skrót": """Przekształć ten materiał edukacyjny w zwięzłe notatki do nauki:

✓ Wyciągnij TYLKO najważniejsze informacje (30% oryginału)
✓ Użyj jasnych nagłówków i podpunktów
✓ Wytłuszcz kluczowe terminy
✓ Dodaj krótkie wyjaśnienia trudnych pojęć
✓ Formatuj w sposób ułatwiający zapamiętywanie

Materiał:""",
                
                "Rozszerzone z wyjaśnieniami": """Przekształć ten materiał edukacyjny w kompleksowe notatki do nauki:

✓ Zachowaj wszystkie ważne informacje
✓ DODAJ proste wyjaśnienia trudnych pojęć (jakbyś tłumaczył koledze)
✓ DODAJ praktyczne przykłady gdzie to możliwe
✓ Użyj analogii dla skomplikowanych tematów
✓ Strukturyzuj: nagłówki → podpunkty → wyjaśnienia
✓ Wytłuszcz najważniejsze terminy
✓ Usuń tylko organizacyjne info (daty sprawdzianów itp.)

Format idealny do nauki! Pisz jasno i przystępnie.

Materiał:"""
            }
            
            # Dodaj custom prompt jeśli jest
            prompt_finalny = prompty[tryb] + "\n\n" + tekst
            if custom_prompt.strip():
                prompt_finalny += f"\n\n⚠️ WAŻNE - zwróć szczególną uwagę na: {custom_prompt}"
            
            # Generuj notatki
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "user", "content": prompt_finalny}
                ]
            )
            notatki = response.choices[0].message.content
            
            # Weryfikacja i poprawa (opcjonalna)
            if weryfikacja:
                with st.spinner("Sprawdzam i poprawiam..."):
                    verify_response = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {"role": "user", "content": f"""Zweryfikuj te notatki edukacyjne:

NOTATKI:
{notatki}

ORYGINAŁ:
{tekst}

ZADANIE:
1. Sprawdź czy są błędy merytoryczne
2. Dodaj wyjaśnienia tam gdzie może być niejasne
3. Upewnij się że format jest przyjazny do nauki
4. Jeśli coś ważnego pominięto - dodaj

Zwróć poprawioną wersję notatek."""}
                        ]
                    )
                    notatki = verify_response.choices[0].message.content
            
            # Zapisz notatki w session_state
            st.session_state['notatki'] = notatki
            
            # Pokaż
            st.success("✅ Gotowe!")
            st.markdown(notatki)
            
            # Word do pobrania
            doc_word = Document()
            doc_word.add_heading('Notatki', 0)
            
            # Dodaj info o trybie
            doc_word.add_paragraph(f"Tryb: {tryb}")
            if custom_prompt.strip():
                doc_word.add_paragraph(f"Dodatkowe instrukcje: {custom_prompt}")
            doc_word.add_paragraph("")  # Pusta linia
            
            doc_word.add_paragraph(notatki)
            
            bio = io.BytesIO()
            doc_word.save(bio)
            
            st.download_button(
                "📥 Pobierz Word",
                bio.getvalue(),
                "notatki.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# ===== TAB 2: QUIZ =====
with tab2:
    st.subheader("🎯 Sprawdź swoją wiedzę!")
    
    if 'pdf_tekst' not in st.session_state:
        st.warning("⚠️ Najpierw wygeneruj notatki w zakładce 'Generuj Notatki'")
    else:
        st.success("✅ Materiał załadowany! Gotowy do quizu.")
        
        # Wybór trybu quizu
        quiz_mode = st.radio(
            "Tryb quizu:",
            ["💬 Chat - zadawaj mi pytania", "❓ Quiz - pytaj mnie", "🎲 Losowe pytania"]
        )
        
        # Inicjalizacja historii chatu
        if 'chat_history' not in st.session_state:
            st.session_state['chat_history'] = []
        
        if 'quiz_pytanie' not in st.session_state:
            st.session_state['quiz_pytanie'] = None
            
        if 'quiz_odpowiedz' not in st.session_state:
            st.session_state['quiz_odpowiedz'] = None
        
        # ===== TRYB 1: CHAT =====
        if quiz_mode == "💬 Chat - zadawaj mi pytania":
            st.markdown("---")
            st.info("💡 Zadaj pytanie o materiał, a ja Ci odpowiem i wytłumaczę!")
            
            # Wyświetl historię
            for msg in st.session_state['chat_history']:
                if msg['role'] == 'user':
                    st.markdown(f"**Ty:** {msg['content']}")
                else:
                    st.markdown(f"**🤖 AI:** {msg['content']}")
                st.markdown("---")
            
            # Input
            user_question = st.text_input("Twoje pytanie:", key="user_chat")
            
            if st.button("Zapytaj") and user_question:
                with st.spinner("Myślę..."):
                    # Dodaj pytanie do historii
                    st.session_state['chat_history'].append({
                        'role': 'user',
                        'content': user_question
                    })
                    
                    # Zapytaj AI
                    response = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {"role": "system", "content": f"Jesteś nauczycielem. Odpowiadaj na pytania ucznia na podstawie tego materiału:\n\n{st.session_state['pdf_tekst']}\n\nWyjaśniaj prosto i przystępnie."},
                            {"role": "user", "content": user_question}
                        ]
                    )
                    
                    answer = response.choices[0].message.content
                    
                    # Dodaj odpowiedź do historii
                    st.session_state['chat_history'].append({
                        'role': 'assistant',
                        'content': answer
                    })
                    
                    st.rerun()
        
        # ===== TRYB 2: QUIZ - AI PYTA =====
        elif quiz_mode == "❓ Quiz - pytaj mnie":
            st.markdown("---")
            st.info("💡 Kliknij 'Następne pytanie' - odpowiedz, a ja sprawdzę Twoją odpowiedź!")
            
            if st.button("📝 Następne pytanie"):
                with st.spinner("Generuję pytanie..."):
                    response = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {"role": "system", "content": f"Jesteś nauczycielem. Zadaj jedno konkretne pytanie sprawdzające wiedzę z tego materiału:\n\n{st.session_state['pdf_tekst']}\n\nPytanie powinno być konkretne i nie za trudne. Zwróć tylko pytanie, bez odpowiedzi."},
                            {"role": "user", "content": "Zadaj mi pytanie"}
                        ]
                    )
                    
                    st.session_state['quiz_pytanie'] = response.choices[0].message.content
                    st.session_state['quiz_odpowiedz'] = None
            
            if st.session_state['quiz_pytanie']:
                st.markdown(f"### ❓ {st.session_state['quiz_pytanie']}")
                
                user_answer = st.text_area("Twoja odpowiedź:", key="quiz_answer")
                
                if st.button("✅ Sprawdź odpowiedź") and user_answer:
                    with st.spinner("Sprawdzam..."):
                        response = client.chat.completions.create(
                            model="gpt-4o-mini",
                            messages=[
                                {"role": "system", "content": f"Jesteś nauczycielem sprawdzającym odpowiedź ucznia. Materiał:\n\n{st.session_state['pdf_tekst']}\n\nOceń odpowiedź: czy jest poprawna, co było dobre, co można poprawić. Bądź wyrozumiały ale konkretny."},
                                {"role": "user", "content": f"Pytanie: {st.session_state['quiz_pytanie']}\n\nOdpowiedź ucznia: {user_answer}\n\nOceń tę odpowiedź."}
                            ]
                        )
                        
                        st.session_state['quiz_odpowiedz'] = response.choices[0].message.content
                
                if st.session_state['quiz_odpowiedz']:
                    st.success("📊 Ocena:")
                    st.markdown(st.session_state['quiz_odpowiedz'])
        
        # ===== TRYB 3: LOSOWE PYTANIA =====
        else:
            st.markdown("---")
            st.info("💡 Wygeneruję serię pytań - odpowiadaj kolejno!")
            
            liczba_pytan = st.slider("Ile pytań?", 3, 10, 5)
            
            if st.button("🎲 Rozpocznij quiz"):
                with st.spinner("Generuję pytania..."):
                    response = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {"role": "system", "content": f"Wygeneruj {liczba_pytan} pytań sprawdzających wiedzę z materiału. Zwróć TYLKO JSON:\n{{'pytania': ['pytanie1', 'pytanie2', ...]}}\n\nMateriał:\n\n{st.session_state['pdf_tekst']}"},
                            {"role": "user", "content": "Wygeneruj pytania"}
                        ]
                    )
                    
                    try:
                        pytania_json = json.loads(response.choices[0].message.content)
                        st.session_state['quiz_pytania'] = pytania_json['pytania']
                        st.session_state['quiz_index'] = 0
                        st.session_state['quiz_wyniki'] = []
                    except:
                        st.error("Błąd generowania pytań, spróbuj ponownie")
            
            if 'quiz_pytania' in st.session_state and st.session_state['quiz_index'] < len(st.session_state['quiz_pytania']):
                current = st.session_state['quiz_index']
                st.markdown(f"### Pytanie {current + 1}/{len(st.session_state['quiz_pytania'])}")
                st.markdown(f"**{st.session_state['quiz_pytania'][current]}**")
                
                user_ans = st.text_area("Twoja odpowiedź:", key=f"ans_{current}")
                
                if st.button("Dalej ➡️") and user_ans:
                    with st.spinner("Sprawdzam..."):
                        response = client.chat.completions.create(
                            model="gpt-4o-mini",
                            messages=[
                                {"role": "system", "content": f"Oceń odpowiedź (1-5 pkt). Materiał:\n\n{st.session_state['pdf_tekst']}"},
                                {"role": "user", "content": f"Pytanie: {st.session_state['quiz_pytania'][current]}\nOdpowiedź: {user_ans}\n\nOceń krótko (max 2 zdania) i daj punkty 1-5."}
                            ]
                        )
                        
                        st.session_state['quiz_wyniki'].append({
                            'pytanie': st.session_state['quiz_pytania'][current],
                            'odpowiedz': user_ans,
                            'ocena': response.choices[0].message.content
                        })
                        
                        st.session_state['quiz_index'] += 1
                        st.rerun()
            
            elif 'quiz_wyniki' in st.session_state and st.session_state['quiz_wyniki']:
                st.success("🎉 Quiz zakończony!")
                st.markdown("### 📊 Twoje wyniki:")
                
                for i, wynik in enumerate(st.session_state['quiz_wyniki']):
                    st.markdown(f"**{i+1}. {wynik['pytanie']}**")
                    st.markdown(f"Twoja odpowiedź: {wynik['odpowiedz']}")
                    st.info(wynik['ocena'])
                    st.markdown("---")
                
                if st.button("🔄 Nowy quiz"):
                    del st.session_state['quiz_pytania']
                    del st.session_state['quiz_wyniki']
                    st.rerun()
        
        # Reset chatu
        if st.button("🗑️ Wyczyść historię"):
            st.session_state['chat_history'] = []
            if 'quiz_pytanie' in st.session_state:
                del st.session_state['quiz_pytanie']
            if 'quiz_odpowiedz' in st.session_state:
                del st.session_state['quiz_odpowiedz']
            st.rerun()